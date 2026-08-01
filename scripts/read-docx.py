import docx
import sys
import io

# Force UTF-8 output
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

path = sys.argv[1] if len(sys.argv) > 1 else "C:/Users/phant/Downloads/GTBT dũng.docx"
# Alternative path without diacritics if original not found
import os
if not os.path.exists(path):
    path = "C:/Users/phant/Downloads/GTBT dung.docx"
doc = docx.Document(path)

print("=== TITLE ===")
print(doc.core_properties.title or "N/A")
print()

print("=== PARAGRAPHS ===")
for p in doc.paragraphs:
    if p.text.strip():
        print(p.text)
        print()

print()
print("=== TABLES ===")
for i, table in enumerate(doc.tables):
    print(f"--- Table {i+1} ---")
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        print(" | ".join(cells))
    print()
